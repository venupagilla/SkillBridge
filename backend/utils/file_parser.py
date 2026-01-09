import io
import mimetypes
from fastapi import UploadFile, HTTPException
import pypdf
import docx

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extracts text from an uploaded file (PDF, DOCX, or TXT).
    """
    content_type = file.content_type
    filename = file.filename.lower()
    
    # Read file content
    file_content = await file.read()
    file_stream = io.BytesIO(file_content)
    
    text = ""
    
    try:
        if filename.endswith(".pdf") or content_type == "application/pdf":
            reader = pypdf.PdfReader(file_stream)
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
                
        elif filename.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file_stream)
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Extract tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
                
        elif filename.endswith(".txt") or content_type == "text/plain":
            text = file_content.decode("utf-8")
            
        else:
             # Basic fallback or error
             raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT.")
             
    except Exception as e:
        print(f"Error parsing file: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")
        
    # Reset cursor just in case
    await file.seek(0)
    
    # Post-processing cleanup
    return _clean_text(text)

def _clean_text(text: str) -> str:
    """
    Removes extra whitespace and cleans up the text.
    """
    import re
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n+', '\n', text)
    # Replace multiple spaces with a single space (preserving newlines)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()
